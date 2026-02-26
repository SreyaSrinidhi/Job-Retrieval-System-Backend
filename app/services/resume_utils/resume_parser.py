import io
import re
from typing import Optional
import pdfplumber
import docx
from typing import BinaryIO
from werkzeug.datastructures import FileStorage


# def extract_text_from_pdf(path: str) -> str:
#     parts: list[str] = []
#     with pdfplumber.open(path) as pdf:
#         for page in pdf.pages:
#             # Keeps line breaks reasonably well
#             txt = page.extract_text() or ""
#             parts.append(txt)
#     return "\n\n".join(parts).strip()

def extract_text_from_pdf(stream: BinaryIO) -> str:
  
    parts: list[str] = []
    stream.seek(0)
    with pdfplumber.open(stream) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n\n".join(parts).strip()


# def extract_text_from_docx(path: str) -> str:
#     d = docx.Document(path)
#     # Join paragraphs with newlines to preserve structure
#     return "\n".join(p.text for p in d.paragraphs if p.text is not None).strip()

def extract_text_from_docx(data: bytes) -> str:
    # python-docx expects a path OR a file-like object.
    # Feeding it BytesIO is the standard approach for uploads.
    doc = docx.Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text is not None).strip()

def normalize_resume_text(text: str) -> str:
    # Normalize newlines and whitespace without destroying structure
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Collapse excessive blank lines (keep up to 2)
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Trim trailing spaces on each line
    text = "\n".join(line.rstrip() for line in text.split("\n"))
    return text.strip()

#TODO - take in FileStorage from werkzeug library - no other inputs
def parse_resume_file(upload: FileStorage) -> str:
    # Only input: FileStorage
    filename = (upload.filename or "").lower()

    # Determine extension from filename (most reliable in practice)
    if filename.endswith(".pdf"):
        text = extract_text_from_pdf(upload.stream)

    elif filename.endswith(".docx"):
        # docx parsing is simplest if we read bytes once
        upload.stream.seek(0)
        data = upload.read()
        text = extract_text_from_docx(data)
    else:
        raise ValueError("Unsupported file type (only pdf, docx)")

    return normalize_resume_text(text)


def looks_like_scanned_or_empty(text: str, min_chars: int = 400) -> bool:
    # If extraction yields very little text, likely scanned PDF or extraction failed
    return len(text.strip()) < min_chars